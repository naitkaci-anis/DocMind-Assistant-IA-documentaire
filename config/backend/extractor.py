"""
backend/extractor.py
====================
Extraction de texte depuis des fichiers PDF et DOCX.

Bibliothèques utilisées :
  - PyMuPDF (fitz) : rapide, gère bien les PDFs complexes
  - python-docx    : lit les fichiers Word modernes (.docx)
"""

from pathlib import Path
import csv as _csv
import fitz      # PyMuPDF — pip install pymupdf
import docx      # python-docx — pip install python-docx


def extract_pdf(filepath: str) -> str:
    """
    Extrait tout le texte d'un PDF, page par page.
    Ajoute un marqueur [Page N] pour savoir d'où vient chaque passage.
    """
    doc   = fitz.open(filepath)
    pages = []

    for num, page in enumerate(doc, start=1):
        text = page.get_text().strip()
        if text:
            pages.append(f"[Page {num}]\n{text}")

    doc.close()
    return "\n\n".join(pages)


def extract_docx(filepath: str) -> str:
    """
    Extrait le texte d'un fichier Word.
    Parcourt les paragraphes puis le contenu des tableaux.
    """
    doc   = docx.Document(filepath)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n\n".join(parts)


def extract_txt(filepath: str) -> str:
    """Lit un fichier texte brut (.txt ou .md) en UTF-8."""
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_csv(filepath: str) -> str:
    """
    Lit un CSV et le convertit en texte lisible.
    Chaque ligne devient une phrase : "col1: val1 | col2: val2 ..."
    """
    rows = []
    with open(filepath, "r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = _csv.DictReader(f)
        headers = reader.fieldnames or []
        for i, row in enumerate(reader):
            parts = [f"{h}: {row.get(h, '').strip()}" for h in headers]
            rows.append(f"[Ligne {i+1}] " + " | ".join(parts))
    return "\n".join(rows)


def extract(filepath: str) -> str:
    """
    Détecte automatiquement le type de fichier et extrait son texte.

    Args:
        filepath: chemin absolu vers le fichier
    Returns:
        texte brut extrait
    Raises:
        ValueError: si l'extension n'est pas supportée
    """
    ext = Path(filepath).suffix.lower()

    if ext == ".pdf":
        return extract_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return extract_docx(filepath)
    elif ext in (".txt", ".md"):
        return extract_txt(filepath)
    elif ext == ".csv":
        return extract_csv(filepath)
    else:
        raise ValueError(f"Format non supporté : '{ext}'.")
