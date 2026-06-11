"""
core/loader.py
==============
Extraction du texte brut depuis les fichiers PDF et DOCX.

Bibliothèques utilisées :
  - PyMuPDF (importé comme "fitz") → lecture PDF
  - python-docx                    → lecture Word

Ces deux libs retournent du texte Python classique (str).
"""

import fitz                          # PyMuPDF : pip install pymupdf
from docx import Document            # python-docx : pip install python-docx
from pathlib import Path


def load_pdf(file_path: str | Path) -> str:
    """
    Extrait tout le texte d'un fichier PDF page par page.

    PyMuPDF ouvre le PDF comme un objet itérable de pages.
    Chaque page expose get_text() qui retourne le texte brut.

    Args:
        file_path : chemin vers le fichier .pdf

    Returns:
        Tout le texte du PDF concaténé, pages séparées par une ligne vide.
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"Fichier introuvable : {file_path}")

    pages_text = []

    # fitz.open() charge le PDF en mémoire
    with fitz.open(str(file_path)) as doc:
        for page_num, page in enumerate(doc, start=1):
            # get_text("text") extrait le texte brut (sans mise en forme)
            text = page.get_text("text").strip()

            if text:   # ignore les pages vides (ex: pages de garde image)
                pages_text.append(f"[Page {page_num}]\n{text}")

    if not pages_text:
        raise ValueError(f"Aucun texte extrait de : {file_path.name}. "
                         "Le PDF est peut-être scanné (image). "
                         "Un OCR sera nécessaire (non implémenté ici).")

    return "\n\n".join(pages_text)


def load_docx(file_path: str | Path) -> str:
    """
    Extrait tout le texte d'un fichier Word (.docx).

    python-docx modélise le document comme une liste de "paragraphes".
    Un paragraphe = un bloc de texte séparé par Entrée dans Word.
    Les tableaux sont traités séparément (sinon leur contenu serait perdu).

    Args:
        file_path : chemin vers le fichier .docx

    Returns:
        Tout le texte du document, paragraphes séparés par \n.
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"Fichier introuvable : {file_path}")

    doc = Document(str(file_path))
    lines = []

    # ── Paragraphes normaux ───────────────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # ── Contenu des tableaux ──────────────────────────────────────────────────
    # python-docx expose les tableaux via doc.tables
    # Chaque table → lignes → cellules
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                lines.append(row_text)

    if not lines:
        raise ValueError(f"Aucun texte extrait de : {file_path.name}.")

    return "\n".join(lines)


def load_document(file_path: str | Path) -> str:
    """
    Point d'entrée universel : détecte l'extension et appelle
    le bon loader (PDF ou DOCX).

    Args:
        file_path : chemin vers le fichier

    Returns:
        Texte extrait sous forme de str

    Raises:
        ValueError si l'extension n'est pas supportée
    """
    file_path = Path(file_path)
    ext = file_path.suffix.lower()

    loaders = {
        ".pdf":  load_pdf,
        ".docx": load_docx,
    }

    if ext not in loaders:
        raise ValueError(
            f"Extension '{ext}' non supportée. "
            f"Formats acceptés : {', '.join(loaders.keys())}"
        )

    return loaders[ext](file_path)


# ── Exécution directe pour tester ─────────────────────────────────────────────
# Ce bloc ne s'exécute que si tu lances `python core/loader.py` directement.
# Il est ignoré quand loader.py est importé depuis un autre module.
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage : python core/loader.py <chemin_fichier>")
        sys.exit(1)

    path = sys.argv[1]
    try:
        texte = load_document(path)
        print(f"✓ Extraction réussie ({len(texte)} caractères)")
        print("─" * 60)
        print(texte[:500])   # affiche les 500 premiers caractères
    except Exception as e:
        print(f"✗ Erreur : {e}")
